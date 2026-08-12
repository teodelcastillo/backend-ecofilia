import logging
import os
from dataclasses import dataclass, field
from typing import Callable, Dict

import re

from apps.document.utils.tables import (
    TABLE_BLOCK_RE,
    block_is_inside_tables,
    extract_page_tables,
)

logger = logging.getLogger(__name__)


@dataclass
class ParseResult:
    """Extracted text plus the coverage stats needed to judge completeness.

    ``page_count`` / ``pages_with_text`` are only meaningful for PDFs; other
    formats leave them at 0 and callers skip the coverage gate.
    """

    text: str = ""
    page_count: int = 0
    pages_with_text: int = 0
    image_only_pages: int = 0
    parser: str = ""
    tables_found: int = 0
    pages_with_tables: int = 0

    @property
    def page_coverage(self) -> float:
        if not self.page_count:
            return 1.0
        return self.pages_with_text / self.page_count

    @property
    def chars_per_page(self) -> float:
        if not self.page_count:
            return float(len(self.text))
        return len(self.text) / self.page_count


# A page yielding less than this is treated as having no usable text: scanned
# pages often carry a few characters (a folio number, a vectorized header)
# without any of the actual body content.
MIN_CHARS_PER_PAGE = 40


def _clean_prose_spacing(text: str) -> str:
    # Collapse single newlines to spaces, but PRESERVE them when the following
    # line starts with a list marker (-, *, •, or digit + . / )) or a page
    # marker (<<<PAGE:N>>>) so structural information survives extraction.
    _LIST_START = r'[ \t]*(?:[-*•]|\d+[.)]) '
    _PAGE_MARKER = r'<<<PAGE:\d+>>>'
    text = re.sub(
        rf'(?<!\n)\n(?!\n)(?!{_LIST_START})(?!{_PAGE_MARKER})',
        ' ',
        text,
    )
    # Collapse multiple spaces / tabs
    text = re.sub(r'[ \t]+', ' ', text)
    # Collapse 3+ newlines to 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def clean_text_spacing(text: str) -> str:
    """Normaliza espaciado sin tocar el interior de los bloques de tabla.

    En una tabla linealizada cada salto de línea separa una fila, así que
    colapsarlos convertiría la tabla en una tira ilegible. Se limpia solo la
    prosa y las tablas quedan como párrafos propios.
    """
    source = text or ""
    cleaned: list[str] = []
    cursor = 0
    for match in TABLE_BLOCK_RE.finditer(source):
        prose = _clean_prose_spacing(source[cursor:match.start()])
        if prose:
            cleaned.append(prose)
        cleaned.append(match.group(0).strip())
        cursor = match.end()

    tail = _clean_prose_spacing(source[cursor:])
    if tail:
        cleaned.append(tail)

    return "\n\n".join(cleaned).strip()


# Lazy imports inside helpers so you don't pay the cost unless needed.

HEADER_FOOTER_PCT = 0.05  # top/bottom band to drop on PDFs


def _read_txt(path: str, encoding: str = "utf-8") -> str:
    with open(path, "r", encoding=encoding, errors="replace") as f:
        return f.read()


_PAGE_MARKER_PREFIX = "<<<PAGE:"
_PAGE_MARKER_SUFFIX = ">>>"


def _make_page_marker(page_num: int) -> str:
    return f"{_PAGE_MARKER_PREFIX}{page_num}{_PAGE_MARKER_SUFFIX}"


def _read_pdf_pymupdf(path: str, top_pct: float = HEADER_FOOTER_PCT, bottom_pct: float = HEADER_FOOTER_PCT) -> ParseResult:
    """Extract text blocks from a PDF using PyMuPDF, removing headers/footers by position.

    Each page's content is prefixed with a <<<PAGE:N>>> marker so the chunker can
    track which source page each chunk came from.

    Also counts how many pages actually yielded text and how many carry only
    images, so the caller can tell "fully extracted" from "silently partial".
    """
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    try:
        pages_out = []
        pages_with_text = 0
        image_only_pages = 0
        tables_found = 0
        pages_with_tables = 0

        for page_num, page in enumerate(doc, start=1):
            height = page.rect.height
            top_y = height * top_pct
            bottom_y = height * (1 - bottom_pct)

            # Tablas primero: se linealizan aparte y su región se excluye del
            # texto corrido para no emitir el mismo contenido dos veces, una
            # como filas y otra como cifras sueltas.
            table_blocks, table_bboxes = extract_page_tables(page)
            if table_blocks:
                tables_found += len(table_blocks)
                pages_with_tables += 1

            # Each block: (x0, y0, x1, y1, text, block_no, block_type, ...)
            blocks_out = []
            for x0, y0, x1, y1, text, *_ in page.get_text("blocks"):
                if y1 < top_y or y0 > bottom_y:
                    continue
                if table_bboxes and block_is_inside_tables((x0, y0, x1, y1), table_bboxes):
                    continue
                text = (text or "").strip()
                if text:
                    blocks_out.append(text)

            blocks_out.extend(table_blocks)
            page_text = "\n\n".join(blocks_out)
            if len(page_text.strip()) >= MIN_CHARS_PER_PAGE:
                pages_with_text += 1
            elif page.get_images(full=True):
                # Scanned/graphic page: nothing to extract without OCR.
                image_only_pages += 1

            if blocks_out:
                pages_out.append(_make_page_marker(page_num) + "\n\n" + page_text)

        if tables_found:
            logger.info(
                "PyMuPDF: %d tablas linealizadas en %d páginas de %s.",
                tables_found, pages_with_tables, path,
            )

        return ParseResult(
            text="\n\n".join(pages_out).strip(),
            page_count=doc.page_count,
            pages_with_text=pages_with_text,
            image_only_pages=image_only_pages,
            parser="pymupdf",
            tables_found=tables_found,
            pages_with_tables=pages_with_tables,
        )
    finally:
        doc.close()


def _read_pdf_pypdf2(path: str) -> ParseResult:
    """Fallback PDF reader when PyMuPDF yields nothing or isn't suitable.

    Emits no <<<PAGE:N>>> markers, so chunks parsed this way have no page
    number — the ``parser`` field records which path produced the text.
    """
    import PyPDF2

    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        pages: list[str] = []
        pages_with_text = 0
        for p in reader.pages:
            try:
                page_text = (p.extract_text() or "").strip()
            except Exception:
                page_text = ""
            if len(page_text) >= MIN_CHARS_PER_PAGE:
                pages_with_text += 1
            pages.append(page_text)

        return ParseResult(
            text="\n".join(pages).strip(),
            page_count=len(reader.pages),
            pages_with_text=pages_with_text,
            parser="pypdf2",
        )


def _read_docx(path: str) -> str:
    """Extract text from a DOCX file including tables.

    Detects explicit page breaks (w:br type="page") to emit <<<PAGE:N>>> markers
    so the chunker can tag each chunk with its source page number.
    """
    import docx  # python-docx
    from docx.oxml.ns import qn

    doc = docx.Document(path)
    parts: list[str] = []
    current_page = 1

    # Start with a page marker for page 1
    parts.append(_make_page_marker(current_page))

    # --- Body paragraphs ---
    for p in doc.paragraphs:
        # Detect explicit page breaks inside paragraph runs
        for run in p.runs:
            for br in run._r.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    current_page += 1
                    parts.append(_make_page_marker(current_page))

        if p.text.strip():
            parts.append(p.text.strip())

    # --- Tables → pipe-delimited rows ---
    # python-docx repeats the value of merged cells for every position they occupy;
    # deduplicate adjacent identical values to avoid noise.
    for table in doc.tables:
        table_lines: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            deduped: list[str] = [cells[0]] if cells else []
            for cell in cells[1:]:
                if cell != deduped[-1]:
                    deduped.append(cell)
            row_text = " | ".join(deduped)
            if row_text.strip(" |"):
                table_lines.append(row_text)
        if table_lines:
            parts.append("\n".join(table_lines))

    return "\n\n".join(parts).strip()


def _parse_file(file_path: str) -> ParseResult:
    """
    Parse a file and return its main textual content plus coverage stats.

    Supported:
      - .txt     (UTF-8, with 'replace' for bad bytes)
      - .pdf     (PyMuPDF for layout-aware extraction + header/footer removal; PyPDF2 fallback)
      - .doc/.docx (python-docx)
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    readers: Dict[str, Callable[[str], str]] = {
        ".txt": _read_txt,
        ".doc": _read_docx,
        ".docx": _read_docx,
    }

    if ext == ".pdf":
        try:
            result = _read_pdf_pymupdf(file_path)
            if result.text:
                return result
            logger.warning(
                "PyMuPDF extracted no text from %s (%d pages, %d image-only); "
                "trying PyPDF2.",
                file_path, result.page_count, result.image_only_pages,
            )
        except Exception:
            # PyMuPDF missing or broken. This used to be a silent `pass`, which
            # hid a production-wide fallback to PyPDF2 (no page markers, far
            # worse extraction) for months — never downgrade quietly again.
            logger.exception(
                "PyMuPDF unavailable or failed on %s; falling back to PyPDF2 "
                "(no page numbers, degraded extraction).",
                file_path,
            )
        return _read_pdf_pypdf2(file_path)

    if ext in readers:
        return ParseResult(text=readers[ext](file_path), parser=ext.lstrip("."))

    raise ValueError(f"Unsupported file type: {ext}")


def parse_file_detailed(file_path: str) -> ParseResult:
    """Parse a file, clean its text spacing, and keep the coverage stats."""
    result = _parse_file(file_path)
    result.text = clean_text_spacing(result.text)
    return result


def parse_file(file_path: str) -> str:
    """Wrapper to parse a file and clean its text spacing."""
    return parse_file_detailed(file_path).text